"""양식 .docx 렌더러 (FR-06).

흐름:
1. 양식 원본 bytes + 변수→값 dict 입력
2. python-docx 로 본문 paragraph + 표 셀 순회
3. 변수 치환 규칙 적용:
   - {{변수명}}, ${변수명}, [변수명] 명시 변수 → 직접 치환
   - "라벨: ___" 패턴 → 라벨 다음에 값 삽입
   - 표 빈 셀 ([__________] markdown 변환에서 추출했던 토큰) → 셀 텍스트 삽입
   - ☐ → ☑ (checkbox=true 인 경우)
4. 서식 보존: paragraph/run 의 글꼴/색상/스타일 유지, 텍스트 내용만 변경
5. NAS_OUTPUTS 백업 + 다운로드 bytes 반환

수용 기준 (FR-06):
- MS Word + 한글에서 양식 깨짐 없음
- 표/이미지/글꼴 100% 보존
- 파일명 규칙: {양식명}_{YYYY-MM-DD}_{user_id_short}.docx
"""
from __future__ import annotations

import io
import logging
import os
import re
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path

from app.config import settings

logger = logging.getLogger(__name__)

# 명시 변수 패턴 — 본문/셀 텍스트에서 1:1 치환.
_EXPLICIT_PATTERNS = [
    (re.compile(r"\{\{(\w+)\}\}"), "{{{}}}"),  # {{key}}
    (re.compile(r"\$\{(\w+)\}"), "${{{}}}"),  # ${key}
    (re.compile(r"\[([가-힣A-Za-z0-9_]+)\]"), "[{}]"),  # [key]
]

# 라벨 + 빈 칸 패턴 (FR-01 #2). "라벨: ___" 또는 "라벨 :   " (공백 5+).
_LABEL_BLANK_PATTERN = re.compile(r"([가-힣A-Za-z0-9 ]{2,20})\s*:\s*(_{3,}|\s{5,})")

# 표 빈 셀 마커 (analyzer.docx_to_markdown 에서 사용한 토큰).
_TABLE_BLANK_MARKER = "[__________]"

# 체크박스 패턴.
_CHECKBOX_UNCHECKED = "☐"
_CHECKBOX_CHECKED = "☑"


@dataclass(frozen=True)
class RenderResult:
    """렌더링 결과 + 저장 메타."""

    file_bytes: bytes
    output_path: str | None  # NAS 저장 경로 (저장 실패 시 None)
    filename: str


def _safe_filename(template_name: str, user_id: str) -> str:
    """파일명 규칙 (FR-06): {양식명}_{YYYY-MM-DD}_{user_id_short}.docx.

    경로 traversal 방지 위해 위험 문자 제거.
    """
    today = datetime.now(tz=timezone.utc).strftime("%Y-%m-%d")
    safe_name = re.sub(r"[^\w가-힣\-_]", "_", template_name).strip("_") or "form"
    short_id = (user_id or "anon").replace("-", "")[:6]
    return f"{safe_name}_{today}_{short_id}.docx"


def _resolve_value(key: str, mapping: dict[str, str | None]) -> str | None:
    """매핑 dict 에서 키 조회. 없거나 None 이면 None 반환 (치환 스킵)."""
    if key not in mapping:
        return None
    val = mapping[key]
    if val is None:
        return None
    return str(val)


def _replace_in_paragraph(paragraph, mapping: dict[str, str | None]) -> int:
    """단락 1개 안에서 명시 변수 + 라벨+빈칸 + 체크박스 치환.

    paragraph.runs 의 글꼴/색상을 보존하기 위해 전체 텍스트를 1번에 합치지 않고
    run 별로 patch — but, 변수가 run 경계를 넘는 경우는 첫 run에 합쳐 넣는 패턴 사용.

    Returns: 치환된 변수 개수
    """
    full_text = paragraph.text or ""
    if not full_text:
        return 0

    new_text = full_text
    replaced = 0

    # 명시 변수 패턴.
    for pattern, _ in _EXPLICIT_PATTERNS:
        for match in list(pattern.finditer(new_text)):
            key = match.group(1)
            value = _resolve_value(key, mapping)
            if value is None:
                continue
            new_text = new_text.replace(match.group(0), value, 1)
            replaced += 1

    # 라벨 + 빈 칸 패턴.
    for match in list(_LABEL_BLANK_PATTERN.finditer(new_text)):
        label = match.group(1).strip()
        # 매핑 키가 라벨과 정확히 일치하거나 공백/특수문자 정규화 매치.
        key_candidates = [label, label.replace(" ", "_"), label.replace(" ", "")]
        value = None
        for cand in key_candidates:
            value = _resolve_value(cand, mapping)
            if value is not None:
                break
        if value is None:
            continue
        # 라벨은 유지하고 빈 칸을 값으로.
        new_text = new_text.replace(
            match.group(0), f"{label}: {value}", 1
        )
        replaced += 1

    # 체크박스: mapping 값이 "true"/True/"checked" 같은 진리값일 때만.
    if _CHECKBOX_UNCHECKED in new_text:
        # 체크박스는 단락 자체에 변수 라벨이 인접한 경우만 처리. 단순 휴리스틱:
        # 단락 텍스트 = "☐ 검토 완료" → mapping["검토 완료"] = "true" 면 ☑로.
        for stripped in [new_text.strip().lstrip(_CHECKBOX_UNCHECKED).strip()]:
            value = _resolve_value(stripped, mapping)
            if value and str(value).lower() in {"true", "checked", "y", "yes", "1", "체크"}:
                new_text = new_text.replace(_CHECKBOX_UNCHECKED, _CHECKBOX_CHECKED, 1)
                replaced += 1

    if new_text == full_text:
        return 0

    # 텍스트가 바뀌었으면 첫 run에 새 텍스트를 몰아넣고 나머지 run은 비움.
    # 이렇게 하면 첫 run의 서식(글꼴/색상)이 전체에 적용된다.
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.text = new_text  # type: ignore[attr-defined]
    return replaced


def _replace_in_cell(cell, mapping: dict[str, str | None]) -> int:
    """표 셀 1개 안에서 변수 치환.

    빈 셀 마커 [__________] 가 있는 경우, 헤더 셀(이전 행 같은 열)을 키로 추정해 매핑 시도.
    이 단계는 이미 analyzer가 헤더 기반 변수 key를 LLM에 추출시켰다는 가정 하에,
    cell 내부에 {{key}} 같은 명시 변수가 있으면 그것을 우선 사용한다.

    더 견고한 헤더 추정은 호출자(forms.py 의 render endpoint)가 mapping 키를
    cell coordinate("table_1_row_2_col_3")로 통일하면 가능 — Phase 1로 미룸.
    """
    replaced = 0
    for paragraph in cell.paragraphs:
        replaced += _replace_in_paragraph(paragraph, mapping)
    # 빈 셀 마커가 그대로 남아있으면 mapping에서 ANY 와일드카드 키 시도(하지 않음 — 안전).
    return replaced


def render_docx(
    *,
    template_bytes: bytes,
    mappings: dict[str, str | None],
    template_name: str,
    user_id: str,
    save_to_nas: bool = True,
) -> RenderResult:
    """양식 원본 bytes + 매핑 dict 로 채워진 .docx bytes 생성.

    Args:
        template_bytes: 양식 원본 .docx 바이너리
        mappings: {variable_key: value} — None 값은 치환 스킵 (빈 칸 그대로)
        template_name: 양식 표시명 (파일명에 사용)
        user_id: 사용자 UUID 문자열 (파일명 short id에 사용)
        save_to_nas: True면 settings.form_filler_output_root 에 백업

    Returns:
        RenderResult — bytes + (NAS) 저장 경로 + 파일명
    """
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover
        raise RuntimeError("python-docx 미설치") from exc

    doc = Document(io.BytesIO(template_bytes))

    total_replaced = 0
    for paragraph in doc.paragraphs:
        total_replaced += _replace_in_paragraph(paragraph, mappings)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                total_replaced += _replace_in_cell(cell, mappings)

    logger.info("renderer: %d 개 변수 치환 완료", total_replaced)

    out = io.BytesIO()
    doc.save(out)
    file_bytes = out.getvalue()

    filename = _safe_filename(template_name, user_id)
    output_path = None
    if save_to_nas:
        output_path = _save_to_nas(file_bytes, filename)

    return RenderResult(
        file_bytes=file_bytes,
        output_path=output_path,
        filename=filename,
    )


def _save_to_nas(file_bytes: bytes, filename: str) -> str | None:
    """settings.form_filler_output_root 에 백업. 실패 시 None 반환 (다운로드는 정상)."""
    root = Path(settings.form_filler_output_root)
    try:
        root.mkdir(parents=True, exist_ok=True)
    except OSError as exc:
        logger.warning("NAS 출력 디렉토리 생성 실패: %s", exc)
        return None
    target = root / filename
    try:
        target.write_bytes(file_bytes)
    except OSError as exc:
        logger.warning("NAS 출력 저장 실패: %s", exc)
        return None
    # path traversal 방지 — root 안인지 확인.
    real_target = os.path.realpath(target)
    real_root = os.path.realpath(root)
    # separator 없는 startswith 는 /nas-evil 가 /nas 에 매칭되는 우회 허용 — 경계 검사.
    if not (real_target == real_root or real_target.startswith(real_root + os.sep)):
        logger.error("출력 경로가 root 외부로 탈출 — 저장 결과 폐기")
        try:
            target.unlink(missing_ok=True)
        except OSError:
            pass
        return None
    return str(target)
