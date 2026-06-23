"""양식 분석기 (FR-01).

흐름:
1. .docx bytes → file_hash(SHA256) 계산
2. python-docx 로 markdown 변환 (표 구조 보존)
3. 캐시 룩업 (file_hash 매칭) 은 라우터 책임 — 본 모듈은 호출 시 분석 실행
4. Claude Sonnet 4.6 호출 (system 프롬프트 캐시)
5. JSON 응답 파싱 + 검증 → List[Variable]

수용 기준 (FR-01): 5초 이내, 표 빈 셀 80% 이상 인식, 동일 file_hash → 동일 결과.
"""
from __future__ import annotations

import hashlib
import io
import json
import logging
import re
from dataclasses import dataclass

from app.config import settings
from app.services.form_filler import guardrails, prompts
from app.services.llm.client import LLMResponse, call_claude

logger = logging.getLogger(__name__)

# 명시 변수 패턴 — markdown 변환에서 양식 안에 있던 그대로 보존.
_EXPLICIT_VAR_HINT_CHARS = {"{", "}", "[", "]", "$", "_", "☐", "□"}

# 단락 빈칸(라벨:___ / 연속 밑줄 / 체크박스) 감지 — 좌표 앵커를 붙일 대상.
# renderer._LABEL_BLANK_PATTERN 과 의미를 맞춘다.
_PARA_BLANK_PATTERN = re.compile(r"(_{3,}|\s{5,})")
_PARA_CHECKBOX_CHARS = ("☐", "□")

# analyzer 분류 → 스키마 허용 타입 매핑(스키마 외 값은 text 폴백).
_ANALYZER_TYPE_MAP = {"enum": "text", "checkbox": "boolean", "image": "text"}


def cell_anchor(table_idx: int, row_idx: int, col_idx: int) -> str:
    """표 셀 좌표 → 안정적 앵커 키. renderer 가 동일 좌표로 역치환한다.

    table_idx 는 1-base(분석/렌더 동일), row_idx 는 list(table.rows) 위치(헤더=0),
    col_idx 는 row.cells 위치. analyzer/renderer 가 같은 python-docx 순회를 쓰므로
    동일 .docx 에 대해 좌표가 일치한다(빈 셀을 본문 텍스트 매칭 없이 채울 수 있게 함).
    """
    return f"t{table_idx}r{row_idx}c{col_idx}"


def para_anchor(para_idx: int) -> str:
    """본문 단락 좌표 → 안정적 앵커 키 (doc.paragraphs 위치)."""
    return f"p{para_idx}"


@dataclass(frozen=True)
class FormVariable:
    """추출된 변수 1개. T5-A의 schemas.form_filler.FormVariable과 정합 가정."""

    key: str
    label: str
    type: str
    location: str
    confidence: float
    required: bool
    default: str | None


@dataclass(frozen=True)
class AnalyzeResult:
    """양식 분석 결과 + 비용 메타. 라우터가 form_templates 저장에 사용."""

    file_hash: str
    variables: list[FormVariable]
    raw_markdown: str
    llm_response: LLMResponse


def compute_file_hash(file_bytes: bytes) -> str:
    """.docx 파일의 SHA256 16진 문자열. 캐시 키 (PRD FR-01 수용 기준)."""
    return hashlib.sha256(file_bytes).hexdigest()


def docx_to_markdown(file_bytes: bytes) -> str:
    """.docx 를 LLM 입력용 markdown 으로 변환.

    표는 markdown 표 문법으로, 빈 셀은 명시적으로 `[__________]` 토큰으로 치환해
    Claude가 빈 칸을 인식하기 쉽게 만든다.
    """
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover
        raise RuntimeError("python-docx 미설치") from exc

    doc = Document(io.BytesIO(file_bytes))
    parts: list[str] = []

    # 본문 단락 — 빈칸(라벨:___ / 밑줄 / 체크박스)이 있으면 좌표 앵커 {{p{i}}} 를
    # 그 자리에 넣어 LLM 이 앵커를 변수 key 로 쓰게 한다(renderer 가 좌표로 역치환).
    for pi, para in enumerate(doc.paragraphs):
        text = para.text or ""
        if text.strip():
            parts.append(_anchor_paragraph(text, pi))
        # 빈 단락은 한 줄 띄움으로 보존 — 양식의 시각적 구조를 유지.
        elif parts and parts[-1] != "":
            parts.append("")

    # 표 — markdown 표로 변환. 빈 데이터 셀은 좌표 앵커 {{t..r..c..}} 로 표기.
    for table_idx, table in enumerate(doc.tables, start=1):
        parts.append(f"\n[표 {table_idx}]")
        rows = list(table.rows)
        if not rows:
            continue
        # 헤더 추정: 첫 행(row_idx=0)이 헤더 — 라벨이라 앵커를 붙이지 않는다.
        header_cells = [_cell_text(c) for c in rows[0].cells]
        parts.append("| " + " | ".join(header_cells) + " |")
        parts.append("| " + " | ".join(["---"] * len(header_cells)) + " |")
        for ri, row in enumerate(rows[1:], start=1):
            cells = [
                _cell_text(c, anchor=cell_anchor(table_idx, ri, ci))
                for ci, c in enumerate(row.cells)
            ]
            parts.append("| " + " | ".join(cells) + " |")

    return "\n".join(parts).strip()


def _anchor_paragraph(text: str, para_idx: int) -> str:
    """본문 단락에 빈칸이 있으면 그 자리에 좌표 앵커 {{p{i}}} 를 넣어 반환.

    빈칸이 없으면 원문 그대로. 체크박스(☐)는 라벨 끝에 앵커를 덧붙인다.
    """
    anchor = "{{" + para_anchor(para_idx) + "}}"
    if _PARA_BLANK_PATTERN.search(text):
        # 첫 빈칸(밑줄/공백덩어리)만 앵커로 치환 — 라벨은 그대로 둔다.
        return _PARA_BLANK_PATTERN.sub(anchor, text, count=1).replace("|", "/")
    if any(ch in text for ch in _PARA_CHECKBOX_CHARS):
        return f"{text} {anchor}".replace("|", "/")
    return text.replace("|", "/")


def _cell_text(cell, *, anchor: str | None = None) -> str:
    """표 셀 텍스트. 빈 셀은 좌표 앵커(있으면) 또는 `[__________]` 토큰으로 치환.

    좌표 앵커({{t..r..c..}})를 주면 LLM 이 그것을 변수 key 로 쓰고, renderer 가
    동일 좌표의 빈 셀에 값을 넣는다(헤더 텍스트 재매칭 불필요 — 표 채움 복원).
    """
    text = (cell.text or "").strip()
    if not text:
        return ("{{" + anchor + "}}") if anchor else "[__________]"
    return text.replace("|", "/")  # markdown 표 구분자 충돌 방지


def _coerce_variable(raw: dict, fallback_idx: int) -> FormVariable | None:
    """LLM 응답의 변수 1개를 dataclass 로 정규화. 부분 누락은 합리적 기본값."""
    key = (raw.get("key") or f"var_{fallback_idx}").strip()
    if not key:
        return None
    label = (raw.get("label") or key).strip()
    # analyzer 분류 어휘 → 스키마 허용 타입으로 정규화.
    # 스키마(schemas.form_filler.FormVariableType): text/number/date/boolean/currency/table_row.
    # 과거엔 enum/checkbox/image 를 그대로 내보내 VariableSchema 검증에서 500 위험이 있었다.
    var_type = (raw.get("type") or "text").strip()
    var_type = _ANALYZER_TYPE_MAP.get(var_type, var_type)
    if var_type not in {"text", "number", "date", "boolean", "currency", "table_row"}:
        var_type = "text"
    location = (raw.get("location") or "").strip()
    try:
        confidence = float(raw.get("confidence", 0.5))
    except (TypeError, ValueError):
        confidence = 0.5
    confidence = max(0.0, min(1.0, confidence))
    required = bool(raw.get("required", False))
    default = raw.get("default")
    if default is not None and not isinstance(default, str):
        default = str(default)
    return FormVariable(
        key=key,
        label=label,
        type=var_type,
        location=location,
        confidence=confidence,
        required=required,
        default=default,
    )


def parse_analyze_response(raw_text: str) -> list[FormVariable]:
    """Claude 응답 JSON → FormVariable 리스트.

    파싱 실패 시 ValueError. 라우터가 422로 응답.
    """
    parsed = guardrails.parse_strict_json(raw_text)
    raw_vars = parsed.get("variables")
    if not isinstance(raw_vars, list):
        raise ValueError("응답 JSON에 'variables' 배열이 없음")
    variables: list[FormVariable] = []
    seen_keys: set[str] = set()
    for idx, raw in enumerate(raw_vars):
        if not isinstance(raw, dict):
            continue
        coerced = _coerce_variable(raw, idx)
        if coerced is None:
            continue
        # 중복 key 제거 (FR-01: 동일 변수 여러 등장 → 1개로 통합).
        if coerced.key in seen_keys:
            continue
        seen_keys.add(coerced.key)
        variables.append(coerced)
    # 50개 초과 시 절단 (FR-01 수용 기준).
    cap = settings.form_filler_max_variables
    if len(variables) > cap:
        logger.warning("변수 %d개 검출 → %d개로 절단", len(variables), cap)
        variables = variables[:cap]
    return variables


def analyze_form(file_bytes: bytes, *, job_metadata: dict | None = None) -> AnalyzeResult:
    """양식 .docx bytes를 받아 변수 목록 + 원본 markdown + LLM 응답 메타를 반환.

    호출자(라우터)가 캐시 룩업 후 cache miss인 경우에만 실행할 것.
    """
    file_hash = compute_file_hash(file_bytes)
    markdown = docx_to_markdown(file_bytes)
    if not markdown:
        raise ValueError("양식 markdown 변환 결과가 비어있음 — 빈 .docx 또는 추출 실패")

    system, messages = prompts.render_analyze_messages(markdown)
    llm_response = call_claude(
        system_prompt=system,
        messages=messages,
        model=settings.form_filler_sonnet_model,
        max_tokens=8192,  # 한국어 JSON 응답에서 4096은 자주 잘림 (변수 30+개 양식)
        temperature=0,  # 구조화 추출 — 결정적 출력(값 흔들림·JSON 깨짐 방지)
        cache_system=True,
        cache_user_first=False,  # 양식 본문은 분석 단계에서 1회만 호출 — 캐시 이득 적음.
        trace_name="form_filler.analyze_form",
        trace_metadata={"file_hash": file_hash, **(job_metadata or {})},
    )

    variables = parse_analyze_response(llm_response.text)
    return AnalyzeResult(
        file_hash=file_hash,
        variables=variables,
        raw_markdown=markdown,
        llm_response=llm_response,
    )
