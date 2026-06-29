"""구조화 섹션 → 디자인된 .docx 바이트.

기본 흰 문서에 글만 찍던 기존 렌더 교체. 표지페이지 + 브랜드 제목 스타일 + 목차(TOC) +
스타일 표(헤더 음영/줄무늬) + 불릿 + 페이지번호. body 의 불릿/표/차트는
markdown_blocks.parse_blocks 로 해석한다. 차트는 docx 네이티브 미지원이라 값 표로 폴백.

회사 .docx 템플릿(settings.docgen_docx_template)이 있으면 그 스타일을 베이스로 쓴다.
"""
from __future__ import annotations

import io
from datetime import date

from app.services.docgen.markdown_blocks import Block, parse_blocks
from app.services.docgen.theme import Theme, get_theme

_BULLET_STYLES = ("List Bullet", "List Bullet 2", "List Bullet 3")


def _hex(color: tuple[int, int, int]) -> str:
    return "{:02X}{:02X}{:02X}".format(*color)


def _set_cell_shading(cell, color: tuple[int, int, int]) -> None:
    """표 셀 배경색(w:shd)."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), _hex(color))
    cell._tc.get_or_add_tcPr().append(shd)


def _add_field(paragraph, instr: str, fallback: str) -> None:
    """필드 코드(TOC/PAGE 등) 삽입. Word에서 F9로 갱신. fallback은 미갱신 시 표시."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr_el = OxmlElement("w:instrText")
    instr_el.set(qn("xml:space"), "preserve")
    instr_el.text = instr
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    fb = OxmlElement("w:t")
    fb.text = fallback
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    r = run._r
    r.append(begin)
    r.append(instr_el)
    r.append(sep)
    r.append(fb)
    r.append(end)


def _style_headings(doc, theme: Theme) -> None:
    """Title/Heading 1·2 스타일을 브랜드 색·폰트로."""
    from docx.shared import Pt, RGBColor

    specs = [("Title", 30, theme.primary, True),
             ("Heading 1", 18, theme.primary, True),
             ("Heading 2", 14, theme.accent, True)]
    for name, size, color, bold in specs:
        try:
            st = doc.styles[name]
        except KeyError:
            continue
        st.font.name = theme.heading_font
        st.font.size = Pt(size)
        st.font.bold = bold
        st.font.color.rgb = RGBColor(*color)
    try:
        normal = doc.styles["Normal"]
        normal.font.name = theme.body_font
        normal.font.size = Pt(10.5)
        normal.font.color.rgb = RGBColor(*theme.text)
    except KeyError:
        pass


def _add_footer_page_number(doc, theme: Theme) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    section = doc.sections[0]
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if theme.footer_text:
        p.add_run(f"{theme.footer_text}    ").italic = True
    _add_field(p, "PAGE", "1")


def _cover_page(doc, theme: Theme, title: str, subtitle: str) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    # 로고(있으면 상단).
    if theme.logo_path:
        try:
            from docx.shared import Inches
            lp = doc.add_paragraph()
            lp.add_run().add_picture(theme.logo_path, height=Inches(0.7))
        except Exception:  # noqa: BLE001
            pass
    for _ in range(6):
        doc.add_paragraph()
    t = doc.add_paragraph(style="Title")
    t.alignment = WD_ALIGN_PARAGRAPH.LEFT
    t.add_run(title or "문서")
    # 액센트 구분선(아래 테두리 가진 빈 문단).
    _add_bottom_border(doc.add_paragraph(), theme.accent)
    if subtitle:
        s = doc.add_paragraph()
        run = s.add_run(subtitle)
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(*theme.muted)
    doc.add_page_break()


def _add_bottom_border(paragraph, color: tuple[int, int, int]) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    pPr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "18")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), _hex(color))
    pbdr.append(bottom)
    pPr.append(pbdr)


def _enable_update_fields(doc) -> None:
    """문서 열 때 TOC/PAGE 필드를 자동 갱신하도록 settings에 표시(Word)."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    settings_el = doc.settings.element
    if settings_el.find(qn("w:updateFields")) is None:
        upd = OxmlElement("w:updateFields")
        upd.set(qn("w:val"), "true")
        settings_el.append(upd)


def _add_toc(doc, theme: Theme) -> None:
    from docx.shared import Pt, RGBColor

    h = doc.add_paragraph()
    run = h.add_run("목차")
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.name = theme.heading_font
    run.font.color.rgb = RGBColor(*theme.primary)
    toc = doc.add_paragraph()
    _add_field(toc, 'TOC \\o "1-2" \\h \\z \\u', "목차는 Word에서 F9로 업데이트하세요.")
    doc.add_page_break()


def _add_styled_table(doc, theme: Theme, rows: list[list[str]]) -> None:
    from docx.shared import Pt, RGBColor

    if not rows:
        return
    ncol = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncol)
    table.style = "Table Grid"
    for i, r in enumerate(rows):
        is_header = i == 0
        for j in range(ncol):
            cell = table.cell(i, j)
            cell.text = r[j] if j < len(r) else ""
            if is_header:
                _set_cell_shading(cell, theme.primary)
            elif i % 2 == 0:
                _set_cell_shading(cell, theme.table_stripe)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9.5)
                    run.font.name = theme.body_font
                    run.font.bold = is_header
                    run.font.color.rgb = RGBColor(
                        *(theme.white if is_header else theme.text))
    doc.add_paragraph()


def _add_chart_fallback(doc, theme: Theme, data: dict) -> None:
    """docx 네이티브 차트 미지원 → 제목 + 값 표로 표현."""
    from docx.shared import Pt, RGBColor

    if data.get("title"):
        p = doc.add_paragraph()
        run = p.add_run(f"［도표］ {data['title']}")
        run.font.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(*theme.accent)
    cats = data["categories"]
    header = ["구분"] + [s["name"] for s in data["series"]]
    rows = [header]
    for idx, cat in enumerate(cats):
        row = [cat]
        for s in data["series"]:
            vals = s["values"]
            row.append(str(vals[idx]) if idx < len(vals) else "")
        rows.append(row)
    _add_styled_table(doc, theme, rows)


def _add_bullet(doc, text: str, level: int = 0) -> None:
    """레벨별 불릿 스타일 적용(없으면 • 평문 폴백)."""
    style = _BULLET_STYLES[min(level, len(_BULLET_STYLES) - 1)]
    try:
        doc.add_paragraph(text, style=style)
    except KeyError:
        doc.add_paragraph(f"• {text}")


def _add_layout_fallback(doc, theme: Theme, data: dict) -> None:
    """레이아웃 블록 → docx 평문 폴백(데이터 유실 방지).

    docx 는 KPI 카드·타임라인 같은 시각 레이아웃을 그대로 못 그리므로, 같은 정보를
    소제목/표/불릿으로 평탄화해 내용은 보존한다(.pptx 가 시각 디자인 담당).
    """
    from docx.shared import Pt, RGBColor

    kind = str(data.get("layout") or "")
    title = data.get("title") or ""
    if title:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.font.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(*theme.primary)

    if kind == "kpi":
        rows = [["지표", "값", "비고"]]
        for it in data.get("items") or []:
            rows.append([it.get("label", ""), it.get("value", ""), it.get("caption", "")])
        _add_styled_table(doc, theme, rows)
    elif kind == "compare":
        for col in data.get("columns") or []:
            heading = col.get("heading", "")
            if heading:
                hp = doc.add_paragraph()
                hp.add_run(heading).font.bold = True
            for pt in col.get("points") or []:
                _add_bullet(doc, pt, level=1)
    elif kind == "timeline":
        for m in data.get("milestones") or []:
            label = m.get("label") or ""
            head = f"[{label}] {m.get('title', '')}".strip()
            _add_bullet(doc, head)
            if m.get("detail"):
                _add_bullet(doc, m["detail"], level=1)
    elif kind == "steps":
        for idx, st in enumerate(data.get("steps") or [], 1):
            _add_bullet(doc, f"{idx}. {st.get('title', '')}")
            if st.get("detail"):
                _add_bullet(doc, st["detail"], level=1)
    elif kind == "quote":
        p = doc.add_paragraph()
        run = p.add_run(f"“{data.get('text', '')}”")
        run.font.italic = True
        run.font.size = Pt(13)
        if data.get("attribution"):
            doc.add_paragraph(f"— {data['attribution']}")


def _render_section_body(doc, theme: Theme, blocks: list[Block]) -> None:
    for b in blocks:
        if b.kind == "paragraph":
            for ln in b.lines:
                if ln.strip():
                    doc.add_paragraph(ln.strip())
        elif b.kind == "bullets":
            for level, text in b.items:
                if text.strip():
                    _add_bullet(doc, text.strip(), level)
        elif b.kind == "table":
            _add_styled_table(doc, theme, b.rows)
        elif b.kind == "chart":
            _add_chart_fallback(doc, theme, b.data)
        elif b.kind == "layout":
            _add_layout_fallback(doc, theme, b.data)


def build_docx(
    title: str, sections: list[dict], theme: Theme | None = None
) -> bytes:
    """title + sections([{heading, body}]) → 디자인된 .docx 바이트.

    theme 가 주어지면 그 색·폰트로(디자인 프리셋), 없으면 회사 기본 테마.
    """
    from docx import Document

    theme = theme or get_theme()
    doc = Document(theme.docx_template) if theme.docx_template else Document()
    _style_headings(doc, theme)
    _add_footer_page_number(doc, theme)

    today = date.today().strftime("%Y.%m.%d")
    _cover_page(doc, theme, title, today)

    headings = [(s.get("heading") or "").strip() for s in sections]
    if len([h for h in headings if h]) >= 2:
        _add_toc(doc, theme)

    for s in sections:
        heading = (s.get("heading") or "").strip()
        if heading:
            doc.add_heading(heading, level=1)
        body = (s.get("body") or "").strip()
        if body:
            _render_section_body(doc, theme, parse_blocks(body))

    _enable_update_fields(doc)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
